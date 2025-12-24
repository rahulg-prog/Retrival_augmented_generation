import sys
from pathlib import Path

current_file = Path(__file__).resolve()
project_root = current_file.parent.parent
sys.path.insert(0, str(project_root))

from logger.custom_logger import CustomLogger
from exceptions.custom_exception import DocumentPortalException

from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
import PyPDF2
from docx import Document as DocxDocument

logger = CustomLogger().get_logger()

class FileValidator:
    """
    Single Responsibility: Validates file existence and type.
    """
    
    @staticmethod
    def validate_file(file_path: Path, expected_extension: str) -> None:
        """
        Validate that file exists and has the expected extension.
        
        Args:
            file_path: Path object of the file
            expected_extension: Expected file extension (e.g., '.pdf', '.docx')
            
        Raises:
            DocumentPortalException: If file doesn't exist or has wrong extension
        """
        if not file_path.exists():
            raise DocumentPortalException(f"File not found: {file_path}")
        
        if file_path.suffix.lower() != expected_extension.lower():
            raise DocumentPortalException(
                f"Invalid file type. Expected {expected_extension}, got {file_path.suffix}"
            )


class DocumentIngestion(ABC):
    """
    Abstract base class for document ingestion.
    Open/Closed Principle: Open for extension (new document types), closed for modification.
    Dependency Inversion: Depends on abstraction (logger interface).
    """
    
    def __init__(self, file_path: str, logger_instance: Optional[Any] = None):
        """
        Initialize document ingestion.
        
        Args:
            file_path: Path to the document file
            logger_instance: Optional logger instance (Dependency Injection)
        """
        self.file_path = Path(file_path)
        self.logger = logger_instance or logger
        self._validate_file()
        self.logger.info(f"{self.__class__.__name__} initialized for: {self.file_path}")
    
    @abstractmethod
    def _validate_file(self) -> None:
        """Validate the file. To be implemented by subclasses."""
        pass
    
    @abstractmethod
    def extract_text(self) -> str:
        """
        Extract text content from the document.
        
        Returns:
            Extracted text as string
        """
        pass
    
    @abstractmethod
    def extract_structured(self) -> List[str]:
        """
        Extract text in a structured format (pages, paragraphs, etc.).
        
        Returns:
            List of text segments
        """
        pass
    
    def get_file_info(self) -> Dict[str, Any]:
        """
        Get basic file information.
        
        Returns:
            Dictionary with file metadata
        """
        return {
            'file_name': self.file_path.name,
            'file_path': str(self.file_path),
            'file_size': self.file_path.stat().st_size,
            'file_type': self.file_path.suffix
        }


class PDFIngestion(DocumentIngestion):
    """
    Single Responsibility: Handles PDF file ingestion only.
    Liskov Substitution: Can be used anywhere DocumentIngestion is expected.
    """
    
    SUPPORTED_EXTENSION = '.pdf'
    
    def __init__(self, file_path: str, logger_instance: Optional[Any] = None):
        """
        Initialize PDF ingestion with file path.
        
        Args:
            file_path: Path to the PDF file
            logger_instance: Optional logger instance (Dependency Injection)
        """
        super().__init__(file_path, logger_instance)
    
    def _validate_file(self) -> None:
        """Validate that the file is a PDF."""
        FileValidator.validate_file(self.file_path, self.SUPPORTED_EXTENSION)
    
    def extract_text(self) -> str:
        """
        Extract text content from the PDF file.
        
        Returns:
            Extracted text as string
        """
        try:
            self.logger.info(f"Starting text extraction from: {self.file_path}")
            text_pages = self.extract_structured()
            extracted_text = "\n".join(text_pages)
            self.logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF: {str(e)} in data_ingestion.py")
            raise DocumentPortalException(f"Failed to extract text from PDF: {self.file_path}", e)
    
    def extract_structured(self) -> List[str]:
        """
        Extract text content from the PDF file page by page.
        
        Returns:
            List of strings, each representing text from one page
        """
        try:
            self.logger.info(f"Starting page-by-page text extraction from: {self.file_path}")
            text_pages = []
            
            with open(self.file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                num_pages = len(pdf_reader.pages)
                self.logger.info(f"PDF has {num_pages} pages")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_pages.append(text if text else "")
                    self.logger.debug(f"Extracted text from page {page_num + 1}")
            
            self.logger.info(f"Successfully extracted text from {len(text_pages)} pages")
            return text_pages
            
        except Exception as e:
            self.logger.error(f"Error extracting text by pages from PDF: {str(e)} in data_ingestion.py")
            raise DocumentPortalException(f"Failed to extract text by pages from PDF: {self.file_path}", e)


class WordIngestion(DocumentIngestion):
    """
    Single Responsibility: Handles Word document file ingestion only.
    Liskov Substitution: Can be used anywhere DocumentIngestion is expected.
    """
    
    SUPPORTED_EXTENSION = '.docx'
    
    def __init__(self, file_path: str, logger_instance: Optional[Any] = None):
        """
        Initialize Word ingestion with file path.
        
        Args:
            file_path: Path to the Word document file (.docx)
            logger_instance: Optional logger instance (Dependency Injection)
        """
        super().__init__(file_path, logger_instance)
    
    def _validate_file(self) -> None:
        """Validate that the file is a Word document."""
        FileValidator.validate_file(self.file_path, self.SUPPORTED_EXTENSION)
    
    def extract_text(self) -> str:
        """
        Extract text content from the Word document.
        
        Returns:
            Extracted text as string
        """
        try:
            self.logger.info(f"Starting text extraction from: {self.file_path}")
            doc = DocxDocument(self.file_path)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            extracted_text = "\n".join(text_content)
            self.logger.info(f"Successfully extracted {len(extracted_text)} characters from Word document")
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document: {str(e)} in data_ingestion.py")
            raise DocumentPortalException(f"Failed to extract text from Word document: {self.file_path}", e)
    
    def extract_structured(self) -> List[str]:
        """
        Extract text content from the Word document paragraph by paragraph.
        
        Returns:
            List of strings, each representing a paragraph
        """
        try:
            self.logger.info(f"Starting paragraph-by-paragraph text extraction from: {self.file_path}")
            doc = DocxDocument(self.file_path)
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            self.logger.info(f"Successfully extracted {len(paragraphs)} paragraphs from Word document")
            return paragraphs
            
        except Exception as e:
            self.logger.error(f"Error extracting paragraphs from Word document: {str(e)}")
            raise DocumentPortalException(f"Failed to extract paragraphs from Word document: {self.file_path}", e)
    
    def extract_with_tables(self) -> Dict[str, Any]:
        """
        Extract text and table content separately from the Word document.
        
        Returns:
            Dictionary with 'text' and 'tables' keys
        """
        try:
            self.logger.info(f"Starting text and table extraction from: {self.file_path}")
            doc = DocxDocument(self.file_path)
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            tables_data = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_content.append(row_data)
                tables_data.append(table_content)
            
            result = {
                'text': "\n".join(paragraphs),
                'tables': tables_data
            }
            
            self.logger.info(f"Successfully extracted text and {len(tables_data)} tables from Word document")
            return result
            
        except Exception as e:
            self.logger.error(f"Error extracting text and tables from Word document: {str(e)}")
            raise DocumentPortalException(f"Failed to extract text and tables from Word document: {self.file_path}", e)


class DocumentIngestionFactory:
    """
    Factory pattern for creating document ingestion instances.
    Single Responsibility: Creates appropriate ingestion objects.
    Open/Closed: Easy to extend with new document types.
    """
    
    _ingestion_classes: Dict[str, type] = {
        '.pdf': PDFIngestion,
        '.docx': WordIngestion
    }
    
    @classmethod
    def create_ingestion(cls, file_path: str, logger_instance: Optional[Any] = None) -> DocumentIngestion:
        """
        Create appropriate ingestion instance based on file extension.
        
        Args:
            file_path: Path to the document
            logger_instance: Optional logger instance
            
        Returns:
            Appropriate DocumentIngestion subclass instance
            
        Raises:
            DocumentPortalException: If file type is not supported
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        ingestion_class = cls._ingestion_classes.get(extension)
        if not ingestion_class:
            supported = ', '.join(cls._ingestion_classes.keys())
            raise DocumentPortalException(
                f"Unsupported file type: {extension}. Supported types: {supported}"
            )
        
        return ingestion_class(file_path, logger_instance)
    
    @classmethod
    def register_ingestion_type(cls, extension: str, ingestion_class: type) -> None:
        """
        Register a new ingestion type.
        Allows extending functionality without modifying existing code.
        
        Args:
            extension: File extension (e.g., '.txt')
            ingestion_class: DocumentIngestion subclass
        """
        if not issubclass(ingestion_class, DocumentIngestion):
            raise ValueError("ingestion_class must be a subclass of DocumentIngestion")
        cls._ingestion_classes[extension.lower()] = ingestion_class

if __name__ == "__main__":
    # Example usage
    try:
        from config.settings_loader import load_config
        config = load_config("config/config.yaml")
        
        # Test PDF ingestion
        pdf_ingestion = DocumentIngestionFactory.create_ingestion(config['data_source']['pdf_path'])
        pdf_text = pdf_ingestion.extract_text()
        print(f"PDF File Info: {pdf_ingestion.get_file_info()}")
        print(f"Extracted PDF Text (first 200 chars): {pdf_text[:200]}...\n")
        
        # Test Word ingestion
        word_ingestion = DocumentIngestionFactory.create_ingestion(config['data_source']['docx_path'])
        word_text = word_ingestion.extract_text()
        print(f"Word File Info: {word_ingestion.get_file_info()}")
        print(f"Extracted Word Text (first 200 chars): {word_text[:200]}...")
        
    except DocumentPortalException as e:
        logger.error(f"Document ingestion failed: {str(e)}")
